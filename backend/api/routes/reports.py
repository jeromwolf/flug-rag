"""
보고서 내보내기 API
Markdown 콘텐츠를 PDF 또는 DOCX 파일로 변환하여 다운로드.
"""
import io
import logging
import re
from datetime import datetime

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from docx.oxml import OxmlElement

from auth.dependencies import get_current_user
from auth.models import User
from config.settings import settings

logger = logging.getLogger(__name__)
router = APIRouter(tags=["reports"])


class ReportExportRequest(BaseModel):
    content: str          # Markdown content
    format: str           # "pdf" or "docx"
    title: str = "보고서"  # Report title


def _build_html(content: str, title: str) -> str:
    """Markdown을 HTML로 변환하고 스타일이 적용된 전체 HTML 문서를 반환."""
    import markdown

    body_html = markdown.markdown(
        content,
        extensions=["tables", "fenced_code"],
    )

    now_str = datetime.now().strftime("%Y년 %m월 %d일")
    platform = settings.platform_name

    html = f"""<!DOCTYPE html>
<html lang="ko">
<head>
  <meta charset="UTF-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1.0" />
  <title>{title}</title>
  <style>
    @import url('https://fonts.googleapis.com/css2?family=Noto+Sans+KR:wght@400;700&display=swap');

    @page {{
      size: A4;
      margin: 2cm;
      @top-center {{
        content: "{title}  |  {platform}";
        font-family: 'Noto Sans KR', sans-serif;
        font-size: 9pt;
        color: #666;
      }}
      @bottom-center {{
        content: "생성일: {now_str}    페이지 " counter(page) " / " counter(pages);
        font-family: 'Noto Sans KR', sans-serif;
        font-size: 9pt;
        color: #666;
      }}
    }}

    * {{
      box-sizing: border-box;
    }}

    body {{
      font-family: 'Noto Sans KR', sans-serif;
      font-size: 11pt;
      line-height: 1.7;
      color: #1a1a1a;
      background: #ffffff;
    }}

    .report-header {{
      text-align: center;
      border-bottom: 2px solid #003366;
      padding-bottom: 16px;
      margin-bottom: 24px;
    }}

    .report-header h1 {{
      font-size: 20pt;
      font-weight: 700;
      color: #003366;
      margin: 0 0 4px 0;
    }}

    .report-header .meta {{
      font-size: 9pt;
      color: #555;
    }}

    h1, h2, h3, h4, h5, h6 {{
      font-family: 'Noto Sans KR', sans-serif;
      font-weight: 700;
      color: #003366;
      margin-top: 1.4em;
      margin-bottom: 0.4em;
    }}

    h1 {{ font-size: 16pt; border-bottom: 1px solid #ccd6e0; padding-bottom: 4px; }}
    h2 {{ font-size: 13pt; }}
    h3 {{ font-size: 11pt; color: #004080; }}

    p {{
      margin: 0.5em 0 1em 0;
    }}

    table {{
      width: 100%;
      border-collapse: collapse;
      margin: 1em 0;
      font-size: 10pt;
    }}

    th, td {{
      border: 1px solid #b0c4d8;
      padding: 6px 10px;
      text-align: left;
    }}

    th {{
      background-color: #003366;
      color: #ffffff;
      font-weight: 700;
    }}

    tr:nth-child(even) td {{
      background-color: #f0f4f8;
    }}

    ul, ol {{
      padding-left: 1.5em;
      margin: 0.5em 0 1em 0;
    }}

    li {{
      margin-bottom: 0.3em;
    }}

    code {{
      font-family: 'Courier New', Courier, monospace;
      font-size: 9pt;
      background: #f4f4f4;
      padding: 1px 4px;
      border-radius: 3px;
      color: #c0392b;
    }}

    pre {{
      background: #f4f4f4;
      border: 1px solid #ddd;
      border-radius: 4px;
      padding: 12px;
      overflow-x: auto;
      font-size: 9pt;
    }}

    pre code {{
      background: none;
      padding: 0;
      color: #1a1a1a;
    }}

    blockquote {{
      border-left: 4px solid #003366;
      margin: 1em 0;
      padding: 0.5em 1em;
      color: #555;
      background: #f8fafc;
    }}

    a {{
      color: #003366;
    }}
  </style>
</head>
<body>
  <div class="report-header">
    <h1>{title}</h1>
    <div class="meta">{platform} &nbsp;|&nbsp; 생성일: {now_str}</div>
  </div>
  <div class="report-body">
    {body_html}
  </div>
</body>
</html>"""
    return html


def _convert_to_pdf(content: str, title: str) -> bytes:
    """Markdown → HTML → PDF 바이트 반환."""
    from weasyprint import HTML

    html_str = _build_html(content, title)
    pdf_bytes = HTML(string=html_str).write_pdf()
    return pdf_bytes


class DocxHtmlParser:
    """HTML을 파싱하여 python-docx Document에 서식이 적용된 콘텐츠를 추가하는 파서.

    Markdown → HTML (via `markdown` library) → DOCX 변환 파이프라인에서 사용.
    테이블, 인라인 서식(bold/italic/code), 코드 블록, 인용구, 목록 등을 지원.
    """

    def __init__(self, doc: "Document") -> None:  # noqa: F821
        from html.parser import HTMLParser

        # HTMLParser를 내부적으로 사용 — 상속 대신 컴포지션으로 구현
        self._doc = doc
        self._parser = HTMLParser()
        self._parser.handle_starttag = self._handle_starttag
        self._parser.handle_endtag = self._handle_endtag
        self._parser.handle_data = self._handle_data

        # 상태 추적
        self._tag_stack: list[str] = []
        self._current_para = None  # 현재 작성 중인 paragraph
        self._inline_styles: dict[str, bool] = {
            "bold": False,
            "italic": False,
            "code": False,
        }

        # 테이블 상태
        self._in_table = False
        self._table_rows: list[list[str]] = []
        self._current_row: list[str] = []
        self._current_cell_text = ""
        self._header_row_count = 0  # thead 안의 행 수
        self._in_thead = False

        # 코드 블록 상태
        self._in_pre = False
        self._pre_text = ""

        # 인용구 상태
        self._in_blockquote = False

        # 리스트 상태
        self._list_stack: list[str] = []  # "ul" or "ol"
        self._ol_counters: list[int] = []  # ol 카운터

    def feed(self, html: str) -> None:
        """HTML 문자열을 파싱하여 Document에 추가."""
        self._parser.feed(html)

    def _handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        tag = tag.lower()
        self._tag_stack.append(tag)

        if tag == "table":
            self._in_table = True
            self._table_rows = []
            self._header_row_count = 0

        elif tag == "thead":
            self._in_thead = True

        elif tag == "tbody":
            pass  # tbody 자체는 특별 처리 불필요

        elif tag == "tr":
            self._current_row = []

        elif tag in ("th", "td"):
            self._current_cell_text = ""

        elif tag == "pre":
            self._in_pre = True
            self._pre_text = ""

        elif tag == "code" and not self._in_pre:
            self._inline_styles["code"] = True

        elif tag == "strong" or tag == "b":
            self._inline_styles["bold"] = True

        elif tag == "em" or tag == "i":
            self._inline_styles["italic"] = True

        elif tag == "blockquote":
            self._in_blockquote = True

        elif tag == "ul":
            self._list_stack.append("ul")

        elif tag == "ol":
            self._list_stack.append("ol")
            self._ol_counters.append(0)

        elif tag == "li":
            if self._list_stack:
                list_type = self._list_stack[-1]
                if list_type == "ol":
                    self._ol_counters[-1] += 1
            self._current_para = None  # li 시작 시 새 paragraph

        elif tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = int(tag[1])
            self._current_para = self._doc.add_heading("", level=level)

        elif tag == "p":
            if not self._in_table and not self._in_pre:
                if self._in_blockquote:
                    self._current_para = self._doc.add_paragraph()
                    self._apply_blockquote_style(self._current_para)
                else:
                    self._current_para = self._doc.add_paragraph()

        elif tag == "hr":
            self._add_horizontal_rule()

        elif tag == "br":
            if self._current_para is not None:
                self._current_para.add_run("\n")

    def _handle_endtag(self, tag: str) -> None:
        tag = tag.lower()

        if tag == "table":
            self._in_table = False
            self._flush_table()

        elif tag == "thead":
            self._in_thead = False

        elif tag == "tr":
            self._table_rows.append(self._current_row)
            if self._in_thead:
                self._header_row_count += 1
            self._current_row = []

        elif tag in ("th", "td"):
            self._current_row.append(self._current_cell_text.strip())
            self._current_cell_text = ""

        elif tag == "pre":
            self._in_pre = False
            self._flush_code_block()

        elif tag == "code" and not self._in_pre:
            self._inline_styles["code"] = False

        elif tag in ("strong", "b"):
            self._inline_styles["bold"] = False

        elif tag in ("em", "i"):
            self._inline_styles["italic"] = False

        elif tag == "blockquote":
            self._in_blockquote = False

        elif tag == "ul":
            if self._list_stack:
                self._list_stack.pop()

        elif tag == "ol":
            if self._list_stack:
                self._list_stack.pop()
            if self._ol_counters:
                self._ol_counters.pop()

        elif tag == "li":
            self._current_para = None

        elif tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            if self._current_para is not None:
                level = int(tag[1])
                sizes = {1: 16, 2: 13, 3: 11, 4: 11, 5: 10, 6: 10}
                colors = {
                    1: (0, 51, 102),
                    2: (0, 51, 102),
                    3: (0, 64, 128),
                    4: (0, 64, 128),
                    5: (0, 64, 128),
                    6: (0, 64, 128),
                }
                from docx.shared import Pt, RGBColor

                _set_heading_font(
                    self._current_para,
                    Pt(sizes.get(level, 11)),
                    RGBColor(*colors.get(level, (0, 51, 102))),
                )
            self._current_para = None

        elif tag == "p":
            self._current_para = None

        # 태그 스택에서 제거
        if self._tag_stack and self._tag_stack[-1] == tag:
            self._tag_stack.pop()

    def _handle_data(self, data: str) -> None:
        # 테이블 셀 안의 텍스트
        if self._in_table and self._tag_stack and self._tag_stack[-1] in ("th", "td"):
            self._current_cell_text += data
            return

        # <pre><code> 안의 텍스트
        if self._in_pre:
            self._pre_text += data
            return

        # 빈 텍스트 무시
        text = data
        if not text.strip():
            # 공백만 있는 경우 — 인라인 요소 사이 공백 유지
            if self._current_para is not None and text == " ":
                self._current_para.add_run(" ")
            return

        # 리스트 아이템 처리
        if self._list_stack and self._current_para is None:
            list_type = self._list_stack[-1]
            if list_type == "ul":
                self._current_para = self._doc.add_paragraph(style="List Bullet")
            else:
                self._current_para = self._doc.add_paragraph(style="List Number")

        # 일반 텍스트 — 현재 paragraph가 없으면 생성
        if self._current_para is None:
            if self._in_blockquote:
                self._current_para = self._doc.add_paragraph()
                self._apply_blockquote_style(self._current_para)
            else:
                self._current_para = self._doc.add_paragraph()

        # Run 추가 (인라인 서식 적용)
        self._add_formatted_run(self._current_para, text)

    def _add_formatted_run(self, para, text: str) -> None:
        """현재 인라인 스타일을 적용하여 Run을 추가."""
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor

        run = para.add_run(text)
        run.font.name = "맑은 고딕"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

        if self._inline_styles["bold"]:
            run.bold = True

        if self._inline_styles["italic"]:
            run.italic = True

        if self._inline_styles["code"]:
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(192, 57, 43)  # #c0392b
            # 인라인 코드 배경색 (회색)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:val"), "clear")
            shading.set(qn("w:color"), "auto")
            shading.set(qn("w:fill"), "F4F4F4")
            run._element.rPr.append(shading)
        else:
            run.font.size = Pt(11)

    def _flush_table(self) -> None:
        """수집된 테이블 데이터를 Document에 테이블로 추가."""
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor

        if not self._table_rows:
            return

        # 최대 열 수 계산
        max_cols = max(len(row) for row in self._table_rows) if self._table_rows else 0
        if max_cols == 0:
            return

        num_rows = len(self._table_rows)
        table = self._doc.add_table(rows=num_rows, cols=max_cols)
        table.style = "Table Grid"

        # 헤더 행 수 (thead가 없으면 첫 행을 헤더로)
        header_count = self._header_row_count if self._header_row_count > 0 else 1

        for row_idx, row_data in enumerate(self._table_rows):
            row = table.rows[row_idx]
            for col_idx in range(max_cols):
                cell = row.cells[col_idx]
                cell_text = row_data[col_idx] if col_idx < len(row_data) else ""

                # 셀 내용 설정
                cell.text = ""
                para = cell.paragraphs[0]
                run = para.add_run(cell_text)
                run.font.name = "맑은 고딕"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
                run.font.size = Pt(10)

                # 헤더 행 스타일
                if row_idx < header_count:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    # 네이비 배경 (#003366)
                    self._set_cell_shading(cell, "003366")
                elif row_idx % 2 == 0:
                    # 짝수 행 — 제브라 스트라이핑 (#f0f4f8)
                    self._set_cell_shading(cell, "F0F4F8")

                # 셀 패딩
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                tcMar = OxmlElement("w:tcMar")
                for side in ("top", "bottom", "start", "end"):
                    el = OxmlElement(f"w:{side}")
                    el.set(qn("w:w"), "60")
                    el.set(qn("w:type"), "dxa")
                    tcMar.append(el)
                tcPr.append(tcMar)

        # 테이블 테두리 설정
        self._set_table_borders(table)

    def _set_cell_shading(self, cell, color_hex: str) -> None:
        """셀에 배경색 설정."""
        from docx.oxml.ns import qn

        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), color_hex)
        tcPr.append(shading)

    def _set_table_borders(self, table) -> None:
        """테이블에 테두리 스타일 설정."""
        from docx.oxml.ns import qn

        tbl = table._element
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
        borders = OxmlElement("w:tblBorders")

        for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "B0C4D8")
            borders.append(border)

        # 기존 borders 제거 후 새로 추가
        existing = tblPr.find(qn("w:tblBorders"))
        if existing is not None:
            tblPr.remove(existing)
        tblPr.append(borders)

    def _flush_code_block(self) -> None:
        """수집된 pre/code 텍스트를 코드 블록 스타일로 추가."""
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor

        text = self._pre_text.strip("\n")
        if not text:
            return

        para = self._doc.add_paragraph()

        # 코드 블록 배경색 (#f4f4f4)
        pPr = para._element.get_or_add_pPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), "F4F4F4")
        pPr.append(shading)

        # 테두리 (연한 회색)
        pBdr = OxmlElement("w:pBdr")
        for side in ("top", "left", "bottom", "right"):
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:space"), "4")
            border.set(qn("w:color"), "DDDDDD")
            pBdr.append(border)
        pPr.append(pBdr)

        # 들여쓰기 (좌우 0.5cm)
        indent = OxmlElement("w:ind")
        indent.set(qn("w:left"), "284")  # ~0.5cm in twips
        indent.set(qn("w:right"), "284")
        pPr.append(indent)

        run = para.add_run(text)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(26, 26, 26)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

        self._pre_text = ""

    def _apply_blockquote_style(self, para) -> None:
        """인용구 스타일 적용 — 좌측 들여쓰기 + 좌측 테두리."""
        from docx.oxml.ns import qn

        pPr = para._element.get_or_add_pPr()

        # 좌측 들여쓰기 (1.5cm ≈ 851 twips)
        indent = OxmlElement("w:ind")
        indent.set(qn("w:left"), "851")
        pPr.append(indent)

        # 좌측 테두리 (네이비)
        pBdr = OxmlElement("w:pBdr")
        left_border = OxmlElement("w:left")
        left_border.set(qn("w:val"), "single")
        left_border.set(qn("w:sz"), "12")  # 1.5pt
        left_border.set(qn("w:space"), "8")
        left_border.set(qn("w:color"), "003366")
        pBdr.append(left_border)
        pPr.append(pBdr)

        # 배경색 (#f8fafc)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), "F8FAFC")
        pPr.append(shading)

    def _add_horizontal_rule(self) -> None:
        """수평선 추가 — 하단 테두리가 있는 빈 paragraph."""
        from docx.oxml.ns import qn

        para = self._doc.add_paragraph()
        pPr = para._element.get_or_add_pPr()

        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "CCD6E0")
        pBdr.append(bottom)
        pPr.append(pBdr)


def _convert_to_docx(content: str, title: str) -> bytes:
    """Markdown → HTML → DOCX 바이트 반환.

    Markdown을 HTML로 변환한 후, DocxHtmlParser로 파싱하여
    서식이 적용된 DOCX 문서를 생성.
    """
    import markdown
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    doc = Document()

    # 페이지 여백 설정 (2cm)
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 기본 폰트 설정 (Malgun Gothic — 한국어 지원)
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = Pt(18)  # ~1.5 line spacing
    # East Asian font fallback
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    # 제목
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.name = "맑은 고딕"
    title_run.font.color.rgb = RGBColor(0, 51, 102)  # #003366

    # 부제목 (날짜 + 플랫폼명)
    now_str = datetime.now().strftime("%Y년 %m월 %d일")
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run(f"{now_str}  |  {settings.platform_name}")
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.name = "맑은 고딕"
    subtitle_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()  # 빈 줄

    # Markdown → HTML → DOCX 변환
    html_body = markdown.markdown(
        content,
        extensions=["tables", "fenced_code"],
    )

    parser = DocxHtmlParser(doc)
    parser.feed(html_body)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _set_heading_font(para, size, color: "RGBColor") -> None:
    """Heading 단락의 폰트를 한국어 호환으로 설정."""
    from docx.oxml.ns import qn

    for run in para.runs:
        run.font.name = "맑은 고딕"
        run.font.size = size
        run.font.color.rgb = color
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")


def _add_run(para, text: str) -> None:
    """단락에 런을 추가하고 한국어 폰트를 설정."""
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run = para.add_run(text)
    run.font.name = "맑은 고딕"
    run.font.size = Pt(11)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")


@router.post("/export")
async def export_report(
    request: ReportExportRequest,
    current_user: User | None = Depends(get_current_user),
) -> StreamingResponse:
    """Markdown 콘텐츠를 PDF 또는 DOCX 파일로 변환하여 반환."""

    if request.format not in ("pdf", "docx"):
        raise HTTPException(
            status_code=400,
            detail="format은 'pdf' 또는 'docx'여야 합니다.",
        )

    safe_title = re.sub(r'[\\/*?:"<>|]', "_", request.title)

    if request.format == "pdf":
        try:
            pdf_bytes = _convert_to_pdf(request.content, request.title)
        except Exception as exc:
            logger.error("PDF 변환 실패: %s", exc, exc_info=True)
            raise HTTPException(
                status_code=500,
                detail=f"PDF 변환 중 오류가 발생했습니다: {exc}",
            ) from exc

        return StreamingResponse(
            io.BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={
                "Content-Disposition": f'attachment; filename="{safe_title}.pdf"',
            },
        )

    else:  # docx
        try:
            docx_bytes = _convert_to_docx(request.content, request.title)
        except Exception as exc:
            logger.error("DOCX 변환 실패: %s", exc, exc_info=True)
            raise HTTPException(
                status_code=500,
                detail=f"DOCX 변환 중 오류가 발생했습니다: {exc}",
            ) from exc

        return StreamingResponse(
            io.BytesIO(docx_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{safe_title}.docx"',
            },
        )
