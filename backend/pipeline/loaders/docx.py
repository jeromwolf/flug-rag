"""DOCX document loader using python-docx."""

import asyncio
from pathlib import Path

from docx import Document as DocxDocument

from .base import BaseLoader, LoadedDocument


class DocxLoader(BaseLoader):
    """Load DOCX files using python-docx."""

    supported_extensions = [".docx"]

    async def load(self, file_path: str | Path) -> LoadedDocument:
        path = self._validate_file(file_path)
        return await asyncio.to_thread(self._load_sync, path)

    def _load_sync(self, path: Path) -> LoadedDocument:
        doc = DocxDocument(str(path))
        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())

        # Extract tables
        for table in doc.tables:
            table_text = self._extract_table(table)
            if table_text:
                text_parts.append(table_text)

        return LoadedDocument(
            content="\n\n".join(text_parts),
            metadata={
                "filename": path.name,
                "file_type": "docx",
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            },
        )

    def _extract_table(self, table) -> str:
        """Extract table as markdown-like format."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        return "\n".join(rows)
