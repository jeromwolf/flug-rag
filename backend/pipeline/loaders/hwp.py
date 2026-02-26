"""HWP document loader using hwp5txt, olefile, and LibreOffice fallbacks."""

import asyncio
import re
import shutil
import subprocess
import tempfile
from pathlib import Path

from .base import BaseLoader, LoadedDocument

# Shell metacharacters that must not appear in file paths passed to subprocess
_UNSAFE_PATH_CHARS = re.compile(r'[;&|`$(){}!<>\x00]')


class HWPLoader(BaseLoader):
    """Load HWP files. Tries hwp5txt first, then olefile, then LibreOffice."""

    supported_extensions = [".hwp"]

    async def load(self, file_path: str | Path) -> LoadedDocument:
        path = self._validate_file(file_path)
        return await asyncio.to_thread(self._load_sync, path)

    def _load_sync(self, path: Path) -> LoadedDocument:
        # Validate file path does not contain shell metacharacters
        if _UNSAFE_PATH_CHARS.search(str(path)):
            raise ValueError(f"Unsafe characters in file path: {path.name}")

        # Try hwp5txt first (best quality)
        try:
            return self._load_with_hwp5txt(path)
        except Exception:
            pass

        # Try olefile (OLE-based extraction)
        try:
            return self._load_with_olefile(path)
        except Exception:
            pass

        # Fallback: LibreOffice CLI conversion
        try:
            return self._load_with_libreoffice(path)
        except Exception as e:
            raise RuntimeError(
                f"Failed to load HWP file: {path}. "
                f"hwp5txt, olefile, LibreOffice all failed. Error: {e}"
            )

    def _load_with_hwp5txt(self, path: Path) -> LoadedDocument:
        """Extract text using hwp5txt command."""
        with tempfile.TemporaryDirectory() as tmp_dir:
            # Copy file to temp dir to avoid cwd pollution
            tmp_hwp = Path(tmp_dir) / path.name
            shutil.copy2(path, tmp_hwp)

            result = subprocess.run(
                ["hwp5txt", str(tmp_hwp)],
                capture_output=True,
                text=True,
                cwd=tmp_dir,
                timeout=30,
            )

            if result.returncode == 0 and result.stdout.strip():
                text = result.stdout.strip()
                if len(text) > 10:
                    return LoadedDocument(
                        content=text,
                        metadata={
                            "filename": path.name,
                            "file_type": "hwp",
                            "extraction_method": "hwp5txt",
                        },
                    )

            raise RuntimeError("hwp5txt produced no output")

    def _load_with_olefile(self, path: Path) -> LoadedDocument:
        """Load HWP using olefile OLE extraction."""
        import struct

        import olefile

        if not olefile.isOleFile(str(path)):
            raise ValueError("Not a valid OLE file")

        ole = olefile.OleFileIO(str(path))
        text_parts = []

        for stream_name in ole.listdir():
            stream_path = "/".join(stream_name)
            if "bodytext" in stream_path.lower():
                try:
                    data = ole.openstream(stream_name).read()
                    text = self._extract_text_from_bodytext(data, struct)
                    if text.strip():
                        text_parts.append(text.strip())
                except Exception:
                    continue

        ole.close()

        content = "\n\n".join(text_parts)
        if not content.strip():
            raise ValueError("No text extracted via olefile")

        return LoadedDocument(
            content=content,
            metadata={
                "filename": path.name,
                "file_type": "hwp",
                "extraction_method": "olefile",
            },
        )

    @staticmethod
    def _extract_text_from_bodytext(data: bytes, struct) -> str:
        """Extract readable text from HWP bodytext binary data."""
        text_parts = []
        i = 0
        while i < len(data) - 1:
            try:
                char_code = struct.unpack_from("<H", data, i)[0]
                if 0x20 <= char_code < 0xFFFF and char_code not in (0xFEFF, 0xFFFE):
                    char = chr(char_code)
                    if char.isprintable() or char in ("\n", "\r", "\t"):
                        text_parts.append(char)
                    elif char_code < 0x20:
                        if char_code in (0x0A, 0x0D, 0x02):
                            text_parts.append("\n")
                i += 2
            except (struct.error, ValueError):
                i += 2
        return "".join(text_parts)

    def _load_with_libreoffice(self, path: Path) -> LoadedDocument:
        """Convert HWP to DOCX using LibreOffice, then extract text."""
        with tempfile.TemporaryDirectory() as tmp_dir:
            result = subprocess.run(
                [
                    "libreoffice",
                    "--headless",
                    "--convert-to", "docx",
                    "--outdir", tmp_dir,
                    str(path),
                ],
                capture_output=True,
                text=True,
                timeout=60,
            )

            if result.returncode != 0:
                raise RuntimeError(f"LibreOffice conversion failed: {result.stderr}")

            docx_files = list(Path(tmp_dir).glob("*.docx"))
            if not docx_files:
                raise RuntimeError("LibreOffice conversion produced no output")

            from docx import Document as DocxDocument

            doc = DocxDocument(str(docx_files[0]))
            text_parts = [para.text for para in doc.paragraphs if para.text.strip()]

            return LoadedDocument(
                content="\n\n".join(text_parts),
                metadata={
                    "filename": path.name,
                    "file_type": "hwp",
                    "extraction_method": "libreoffice",
                },
            )
